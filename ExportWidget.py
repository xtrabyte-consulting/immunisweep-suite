#!/usr/bin/env python3
"""
Export Widget Module
====================
This module defines the ExportWidget class, which provides a simple GUI for exporting
field sweep data to CSV or DOCX files. The widget allows the user to input evaluation
and criteria data, then save the sweep data using the chosen format.

Dependencies:
    - csv: For writing CSV files.
    - os: For file and directory operations.
    - docx (python-docx): For generating DOCX files.
    - datetime: For timestamping filenames.
    - PyQt5: For creating the widget and dialogs.
"""

import csv
import os
from docx import Document
from datetime import datetime
from PyQt5.QtWidgets import QPushButton, QFileDialog, QVBoxLayout, QHBoxLayout, QWidget, QLabel, QLineEdit, QMessageBox

class ExportWidget(QWidget):
    """
    ExportWidget provides an interface to export field sweep data to CSV or DOCX format.
    
    The widget displays input fields for evaluation and criteria, along with buttons to
    save the data or cancel the export.
    """
    
    def __init__(self, field_strengths = []):
        """
        Initialize the ExportWidget with the given field strength data.
        
        Parameters:
            field_strengths (list, optional): A list of tuples containing field data,
                typically in the form (frequency, field). Defaults to an empty list.
        """
        super().__init__()
        self.field_levels = field_strengths if field_strengths is not None else []
        
        # Create UI elements
        self.evaluation_label = QLabel("Evaluation:", self)
        self.evaluation_input = QLineEdit(self)
        self.criteria_label = QLabel("Criteria:", self)
        self.criteria_input = QLineEdit(self)
        self.csv_button = QPushButton("Save to CSV")
        self.docx_button = QPushButton("Save to DOCX")
        self.cancel_button = QPushButton("Cancel")
        self.label = QLabel("Test Complete. Export sweep data?")
        
        # Set tooltips for buttons
        self.csv_button.setToolTip("Save the data to a CSV file")
        self.docx_button.setToolTip("Save the data to a DOCX file")
                
        # Connect button signals to corresponding slots
        self.csv_button.clicked.connect(self.save_csv)
        self.docx_button.clicked.connect(self.save_docx)
        self.cancel_button.clicked.connect(self.close)

        # Set up layout for the widget
        layout = QVBoxLayout()
        layout.addWidget(self.label)
        layout.addWidget(self.evaluation_label)
        layout.addWidget(self.evaluation_input)
        layout.addWidget(self.criteria_label)
        layout.addWidget(self.criteria_input)
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.csv_button)
        button_layout.addWidget(self.docx_button)
        button_layout.addWidget(self.cancel_button)
        layout.addLayout(button_layout)
        self.setLayout(layout)
        
    def set_output_powers(self, field_levels):
        """
        Set the field level data to be exported.
        
        Parameters:
            field_levels (list): A list of tuples (frequency, field) representing the sweep data.
        """
        self.field_levels = field_levels

    def save_csv(self):
        """
        Open a file dialog to save the field data to a CSV file.
        
        On success, display a success message; otherwise, display an error message.
        """
        
        file_name, _ = QFileDialog.getSaveFileName(self, "Save CSV File", "", "CSV Files (*.csv)")
        if file_name:
            try:
                self.save_to_csv(
                    self.field_levels,
                    self.evaluation_input.text(),
                    self.criteria_input.text(),
                    file_name
                )
                QMessageBox.information(self, "Success", f"Data saved to {file_name}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save file: {e}")
        self.close()

    def save_docx(self):
        """
        Open a file dialog to save the field data to a DOCX file.
        
        On success, display a success message; otherwise, display an error message.
        """
        file_name, _ = QFileDialog.getSaveFileName(self, "Save DOCX File", "", "Word Documents (*.docx)")
        if file_name:
            try:
                self.save_to_docx(self.field_levels, self.evaluation_input.text(), self.criteria_input.text(), file_name)
                QMessageBox.information(self, "Success", f"Text saved to {file_name}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save file: {e}")
        self.close()

    def save_to_csv(self, field_levels, evaluation, criteria, filename=f"output_powers_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.csv"):
        """
        Save the field data to a CSV file.
        
        Parameters:
            field_levels (list): A list of tuples (frequency, field).
            evaluation (str): Evaluation text to be included in each row.
            criteria (str): Criteria text to be included in each row.
            filename (str): The filename for the CSV file.
        """
        with open(filename, mode="w", newline="") as file:
            writer = csv.writer(file)
            writer.writerow(["Level", "Frequency MHz", "Voltage V/m", "Evaluation", "Criteria", "Spec.", "Result"])
            for freq, field in field_levels:
                writer.writerow([self.get_level(field), freq, field, evaluation, criteria, "A", ""])
                
    def save_to_docx(self, field_levels, evaluation, criteria, filename=f"output_powers_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"):
        """
        Save the field data to a DOCX file as a formatted table.
        
        Parameters:
            field_levels (list): A list of tuples (frequency, field).
            evaluation (str): Evaluation text to be included in the table.
            criteria (str): Criteria text to be included in the table.
            filename (str): The filename for the DOCX file.
        """
        doc = Document()
        doc.add_heading("Radiated Immunity Data", level=1)

        # Create a table with rows equal to the number of data points + 1 for the header
        table = doc.add_table(rows=len(field_levels) + 1, cols=7)
        table.style = 'Table Grid'

        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Level"
        hdr_cells[1].text = "Frequency MHz"
        hdr_cells[2].text = "Voltage V/m"
        hdr_cells[3].text = "Evaluation"
        hdr_cells[4].text = "Criteria"
        hdr_cells[5].text = "Spec."
        hdr_cells[6].text = "Result"

        # Add data to the table
        for idx, (freq, field) in enumerate(field_levels):
            row_cells = table.rows[idx + 1].cells
            row_cells[0].text = self.get_level(field)
            row_cells[1].text = str(freq)
            row_cells[2].text = str(field)
            row_cells[3].text = evaluation
            row_cells[4].text = criteria
            row_cells[5].text = 'A'
            row_cells[6].text = ''

        # Save the document
        doc.save(filename)
        
    def get_level(self, field):
        """
        Determine the level category based on the field strength.
        
        Parameters:
            field (float): The measured field strength.
            
        Returns:
            str: A level indicator as a string ("1", "2", or "3").
            Returns "1" if field < 2.5, "2" if 2.5 <= field <= 6, else "3".
        """
        if field < 2.5:
            return "1"
        elif field <= 6 and field >= 2.5:
            return "2"
        else:
            return "3"
